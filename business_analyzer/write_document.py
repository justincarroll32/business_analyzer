import os
import shutil
from docx import Document

from docx.shared import Inches, RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import utilities as ut


class WriteDocument:

    def add_red_line(self, doc: Document) -> None:
        word = 2*"____________________________________________________"
        p = doc.add_paragraph()
        runner = p.add_run(word)
        runner.bold = True
        runner.font.color.rgb = RGBColor(255, 0, 0)

    def add_bold(self, doc: Document, word: str, font_size: int) -> None:
        p = doc.add_paragraph()
        runner = p.add_run(word)
        runner.bold = True
        runner.font.size = Pt(font_size)

    def add_italic(self, doc: Document, word: str, font_size: int) -> None:
        p = doc.add_paragraph()
        runner = p.add_run(word)
        runner.italic = True
        runner.font.size = Pt(font_size)

    def add_regular(self, doc: Document, word: str, font_size: int) -> None:
        p = doc.add_paragraph()
        runner = p.add_run(word)
        runner.font.size = Pt(font_size)

    def get_min_wage_url_from_place(self, places: str) -> str:
        state = (places.split(", ")[1]).replace(" ", "%20")
        return f'https://www.epi.org/minimum-wage-tracker/#/min_wage/{state}'

    def write_to_document(self, doc_name: str, visual: str, radius: float, business_type: str, place: str, county: str, hours: list,
                          ratings: list, servings: list, demographics: tuple) -> None:
        doc = Document()
        # Add a heading
        word = 'Overview and Statistics of Surrounding Businesses'
        p = doc.add_heading()
        runner = p.add_run(word)
        runner.bold = True
        runner.font.color.rgb = RGBColor(0, 0, 0)
        runner.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # HEADER
        self.add_red_line(doc)
        self.add_regular(doc, f'Location: {place}', 10)
        self.add_regular(doc, f'County: {county}', 10)
        self.add_regular(doc, f'Radius: ~{radius} miles', 10)
        self.add_regular(doc, f'Business type: {business_type}', 10)
        self.add_regular(doc, f'Number of businesses found: {ratings[0][0]}', 10)
        min_wage_url = self.get_min_wage_url_from_place(place)
        self.add_regular(doc, f'To track minimum wage for your state: {min_wage_url}', 10)
        self.add_red_line(doc)

        # MAP

        # only adds picture if visual does not catch an error
        if visual != "1":
          self.add_bold(doc, 'Map', 16)
          map_path = visual
          doc.add_picture(map_path, width=Inches(6), height=Inches(4))
          self.add_red_line(doc)

        # HOURS

        word = 'Hours'
        self.add_bold(doc, word, 16)
        stat_names = ['Places that are open the most (in hours)', 'Places that are open the least (in hours)', 'Average hours open per week',
                      'Days that places are closed the most', 'Earliest open places', 'Latest open places']

        for name2, stat in zip(stat_names, hours):
            if stat == 0 or len(stat) == 0:
                self.add_italic(doc, name2, 14)
                doc.add_paragraph('No Data', style='ListBullet')
            else:
                tuple_length = len(stat[0])
                word = name2

                self.add_italic(doc, word, 14)

                table = doc.add_table(rows=0, cols=tuple_length)
                table.autofit = True
                table.style = 'Table Grid'
                if tuple_length == 2:
                    for name, number in stat:
                        row_cells = table.add_row().cells
                        row_cells[0].text = name
                        row_cells[1].text = str(number)
                    doc.add_paragraph()
                else:
                    for day_name, number_days, places in stat:
                        row_cells = table.add_row().cells
                        row_cells[0].text = day_name
                        row_cells[1].text = str(number_days)
                        row_cells[2].text = places
                    doc.add_paragraph()

        self.add_italic(doc, "24/7 places", 14)

        if hours[-1] == []:
            doc.add_paragraph('No places open 24/7', style='ListBullet')
        else:
            for place in hours[-1]:
                doc.add_paragraph(place, style='ListBullet')

        self.add_red_line(doc)
        doc.save(doc_name)

        # RATINGS

        word = 'Ratings:'

        self.add_bold(doc, word, 16)

        ratings_name = ['Average Rating', 'Bad Ratings and Examples', 'Good Ratings and Examples']
        for name, data in zip(ratings_name, ratings):
            if name == 'Average Rating':
                total = data[0]

                word = f'-{name} of {total} Businesses:'

                self.add_italic(doc, word, 14)
                word = str(data[1]) + " stars"
                doc.add_paragraph(word, style='ListBullet')
            elif name == 'Bad Ratings and Examples':
                word = f'-{name} (1-2 stars):'

                self.add_italic(doc, word, 14)

                for item in data:
                    doc.add_paragraph(item, style='ListBullet')
            else:
                word = f'-{name} (4-5 stars):'

                self.add_italic(doc, word, 14)

                for item in data:
                    doc.add_paragraph(item, style='ListBullet')

        self.add_red_line(doc)
        doc.save(doc_name)
        # SERVINGS

        word = 'What Businesses in the Area Serve'

        self.add_bold(doc, word, 16)

        servings_name = ['Serving Percents', 'Suggestions']

        for name, data in zip(servings_name, servings):
            if name == 'Serving Percents':

                word = 'Proportion of what Businesses Offer'

                self.add_italic(doc, word, 14)
                table = doc.add_table(rows=0, cols=2)
                table.autofit = True
                table.style = 'Table Grid'

                for category, percent in data.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = category
                    row_cells[1].text = str(percent)

                doc.add_paragraph()
            else:

                word = 'Suggestions to Offer'

                self.add_italic(doc, word, 14)
                total = data[0]
                suggestions = data[1]
                doc.add_paragraph(f'Less than 50% of businesses have these available:')

                for item in suggestions:
                    doc.add_paragraph(item, style='ListBullet')

        self.add_red_line(doc)
        doc.save(doc_name)
        doc.add_page_break()
        # DEMOGRAPHICS (contains demographics pie charts and migration data)


        self.add_bold(doc, 'Demographics Data', 16)
        self.add_italic(doc, 'Data from United States Census Bureau reflects available data', 8)
        demographics_names = ['Ethnic', 'Age', 'Sex']
        migration_names = ['Domestic', 'International', 'Net']


        if demographics == (404, 404):
            self.add_bold(doc, 'Demographics Data is not available... for now', 14)
        else:
            demographic, migration = demographics
            for name, data in zip(demographics_names, demographic):
                if name == "Ethnic":

                    ethnic_dir = data[0]
                    ethnic_path = data[1]
                    self.add_italic(doc, 'Race', 14)
                    doc.add_picture(ethnic_path, width=Inches(4), height=Inches(3.2))

                    if os.path.exists(ethnic_dir):
                        shutil.rmtree(ethnic_dir)

                elif name == 'Age':

                    age_dir = data[0]
                    age_path = data[1]
                    self.add_italic(doc, 'Age', 14)
                    doc.add_picture(age_path, width=Inches(4), height=Inches(3.2))

                    if os.path.exists(age_dir):
                        shutil.rmtree(age_dir)

                else:

                    doc.add_page_break()
                    sex_dir = data[0]
                    sex_path = data[1]
                    self.add_italic(doc, 'Sex', 14)
                    doc.add_picture(sex_path, width=Inches(4), height=Inches(3.2))
                    if os.path.exists(sex_dir):
                        shutil.rmtree(sex_dir)


            self.add_red_line(doc)
            doc.add_page_break()
            self.add_bold(doc, 'Migration Trends', 16)
            self.add_italic(doc, 'Data from United States Census Bureau reflects available data between 2010-2021', 8)
            for name, data in zip(migration_names, migration):
                if name == 'Domestic':
                    domestic_dir = data[0]
                    domestic_path = data[1]

                    doc.add_picture(domestic_path, width=Inches(6), height=Inches(4))
                    if os.path.exists(domestic_dir):
                        shutil.rmtree(domestic_dir)

                elif name == 'International':
                    international_dir = data[0]
                    international_path = data[1]

                    doc.add_picture(international_path, width=Inches(6), height=Inches(4))
                    if os.path.exists(international_dir):
                        shutil.rmtree(international_dir)

                else:
                    net_dir = data[0]
                    net_path = data[1]

                    doc.add_picture(net_path, width=Inches(6), height=Inches(4))
                    if os.path.exists(net_dir):
                        shutil.rmtree(net_dir)
                        # Add a caption or description

        self.add_red_line(doc)
        doc.save(doc_name)
